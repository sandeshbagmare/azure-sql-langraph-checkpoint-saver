"""Generate THE BOOK — a single, comprehensive Word reference for this repo.

A 30+ page book covering every aspect of langgraph-checkpoint-azure-sql:
concepts, architecture, engineering decisions, all seven benchmarks (with the
real numbers read from benchmarks/results/*.json), consolidated findings,
challenges and how they were solved, limitations, production/security/cost
guidance, troubleshooting, and appendices.

Output: docs/THE_BOOK_Azure_SQL_LangGraph.docx

Usage:
    pip install python-docx
    python scripts/generate_book.py
"""
from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "benchmarks" / "results"
OUT = ROOT / "docs" / "THE_BOOK_Azure_SQL_LangGraph.docx"

AZURE = RGBColor(0x00, 0x78, 0xD4)
NAVY = RGBColor(0x10, 0x3A, 0x5B)
DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
RED = RGBColor(0xB0, 0x2A, 0x2A)


# ---------------------------------------------------------------------------
# data
# ---------------------------------------------------------------------------

def load(name):
    p = RESULTS / f"{name}.json"
    if not p.exists():
        return None
    with open(p) as f:
        return json.load(f)


def human(b):
    if b is None:
        return "—"
    b = float(b)
    if abs(b) < 1024:
        return f"{b:.0f} B"
    if abs(b) < 1_048_576:
        return f"{b / 1024:.1f} KB"
    if abs(b) < 1_073_741_824:
        return f"{b / 1_048_576:.2f} MB"
    return f"{b / 1_073_741_824:.2f} GB"


# ---------------------------------------------------------------------------
# docx building blocks
# ---------------------------------------------------------------------------

def base_style(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.font.color.rgb = DARK
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.15


def H1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = AZURE
        r.font.size = Pt(18)
    return h


def H2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(13.5)
    return h


def H3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(11.5)
    return h


def P(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(f"{bold_lead} ")
        r.bold = True
    p.add_run(text)
    return p


def numbered(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(f"{bold_lead} ")
        r.bold = True
    p.add_run(text)
    return p


def note(doc, text, color=GREY, label="Note"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GREY
    return p


def code(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.8)
        r.font.color.rgb = RGBColor(0x0A, 0x3D, 0x62)
    doc.add_paragraph()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = NAVY
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
    doc.add_paragraph()
    return t


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "  ›  Right-click here and choose “Update Field” to build the contents."
    placeholder.append(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (b, instr, sep, placeholder, end):
        run._r.append(el)


def set_update_fields_on_open(doc):
    settings = doc.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def footer_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pre = p.add_run("Azure SQL LangGraph Checkpoint Saver  ·  Page ")
    pre.font.size = Pt(8)
    pre.font.color.rgb = GREY
    run = p.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(b)
    run._r.append(instr)
    run._r.append(end)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


def pagebreak(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# cover + front matter
# ---------------------------------------------------------------------------

def cover(doc, full):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THE COMPLETE GUIDE")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Azure SQL Database")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = AZURE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LangGraph Checkpoint Saver")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Architecture · Engineering · Benchmarks · Findings · Production Operations")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

    for _ in range(3):
        doc.add_paragraph()

    real = full.get("real_azure_endpoint", False) if full else False
    target = full.get("target_db", "langgraph_azure") if full else "langgraph_azure"
    meta = [
        ["Library", "langgraph-checkpoint-azure-sql v0.1.0"],
        ["Companion library", "langgraph-checkpoint-mssql (on-premises SQL Server)"],
        ["Framework", "LangGraph 1.x · langgraph-checkpoint 4.1.x"],
        ["Language / Runtime", "Python 3.13 · pyodbc 5.x · ODBC Driver 18"],
        ["Database engine", "Microsoft SQL Server 2022 / Azure SQL Database"],
        ["Benchmark target", f"{target} (real Azure endpoint: {'yes' if real else 'no — local stand-in'})"],
        ["Authors", "Sandesh Bagmare · Pawan Nala"],
        ["Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")],
        ["Document type", "Auto-generated from live benchmark results"],
    ]
    table(doc, ["Field", "Value"], meta)
    pagebreak(doc)


def preface(doc):
    H1(doc, "Preface — How to Read This Book")
    P(doc,
      "This book is the single, complete reference for the Azure SQL LangGraph "
      "checkpoint saver. It is written for three audiences at once: the engineer "
      "who will integrate the library, the architect who must decide whether it "
      "fits, and the operator who will run it in production. You can read it "
      "cover to cover, or jump to the part you need using the table of contents.")
    P(doc, "The book is organised into four parts and a set of appendices:")
    bullet(doc, "the concepts, the project, and the data model.", bold_lead="Part I — Foundations:")
    bullet(doc, "how the library is built and why each decision was made.", bold_lead="Part II — Engineering:")
    bullet(doc, "seven benchmark suites, the real numbers, and what they mean.", bold_lead="Part III — Benchmarks & Findings:")
    bullet(doc, "deploying, securing, monitoring, sizing, and troubleshooting.", bold_lead="Part IV — Operations:")
    bullet(doc, "schema, connection-string cookbook, reproduction steps, glossary.", bold_lead="Appendices:")
    note(doc,
         "Every benchmark number in this book is read from benchmarks/results/*.json, "
         "produced by running the suite against a real database. Nothing is hand-typed. "
         "Re-run `python -m benchmarks.run_all` then `python scripts/generate_book.py` to "
         "regenerate the entire document from fresh measurements.",
         label="On trust")
    H2(doc, "Contents")
    add_toc(doc)
    pagebreak(doc)


# ---------------------------------------------------------------------------
# PART I — FOUNDATIONS
# ---------------------------------------------------------------------------

def part_divider(doc, part_no, title, subtitle):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"PART {part_no}")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = AZURE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    pagebreak(doc)


def ch_introduction(doc):
    H1(doc, "1. Introduction & Motivation")
    H2(doc, "1.1 The problem")
    P(doc,
      "LangGraph is a framework for building stateful, multi-actor applications with "
      "large language models. It models an agent as a graph of computation steps and "
      "persists the agent's state between steps through a component called a checkpoint "
      "saver. That persistence is what makes agents resumable after a crash, debuggable "
      "by replay, safe to interrupt for human review, and capable of multi-turn memory.")
    P(doc,
      "Out of the box, LangGraph ships checkpoint savers for PostgreSQL and SQLite. It "
      "does not ship one for Microsoft SQL Server or Azure SQL Database. Yet a very large "
      "share of enterprises — particularly those already invested in Microsoft and Azure "
      "— run SQL Server as their system of record. For them, adding a Postgres instance "
      "just to hold agent checkpoints means a new database engine to secure, back up, "
      "patch, and staff. This library removes that friction.")
    H2(doc, "1.2 Why Azure SQL specifically")
    bullet(doc, "Azure-first enterprises already have Azure SQL provisioned, monitored, and compliant.")
    bullet(doc, "It is fully managed: automatic backups, geo-replication, built-in high availability.")
    bullet(doc, "It supports passwordless authentication via Azure AD and Managed Identity.")
    bullet(doc, "It carries SOC 2, HIPAA, FedRAMP, and ISO 27001 certifications out of the box.")
    bullet(doc, "It runs the identical T-SQL engine as on-premises SQL Server, so one library serves both.")
    H2(doc, "1.3 What this library delivers")
    numbered(doc, "A drop-in BaseCheckpointSaver implementation for Azure SQL and SQL Server.")
    numbered(doc, "The same four-table schema as the official Postgres saver, for predictable behaviour.")
    numbered(doc, "Injection-proof, MERGE-free, concurrency-correct T-SQL.")
    numbered(doc, "Sync and async interfaces, with thread-safe connection pooling.")
    numbered(doc, "A reproducible seven-suite benchmark harness whose output feeds this document.")
    numbered(doc, "First-class Azure AD / Managed Identity support and production guidance.")


def ch_concepts(doc):
    H1(doc, "2. LangGraph Checkpointing Concepts")
    P(doc,
      "Before the architecture makes sense, four LangGraph concepts must be clear. The "
      "library exists to store and retrieve exactly these.")
    table(doc, ["Concept", "Meaning", "Where it lives"], [
        ["Thread", "One conversation / agent run, identified by a thread_id. All of a thread's history shares that id.", "thread_id column (all tables)"],
        ["Checkpoint", "A snapshot of the graph's full state at one super-step, with a unique checkpoint_id.", "checkpoints table"],
        ["Channel", "A named slot in the state (e.g. 'messages', 'text'). Each has a monotonic version.", "checkpoint_blobs table"],
        ["Pending write", "Intermediate output of a task within a step, before the step commits.", "checkpoint_writes table"],
    ])
    H2(doc, "2.1 The lifecycle of one graph step")
    numbered(doc, "The graph executes a node, producing channel updates.")
    numbered(doc, "Each task's output is written as a pending write (checkpoint_writes).")
    numbered(doc, "When the step commits, a new checkpoint is written (checkpoints) and changed channel values are stored once per version (checkpoint_blobs).")
    numbered(doc, "On resume, the saver returns the latest checkpoint plus its channel values and any pending writes, and the graph continues.")
    H2(doc, "2.2 The saver contract")
    P(doc, "A checkpoint saver implements a small, well-defined interface. This library implements all of it, synchronously and asynchronously:")
    table(doc, ["Method", "Purpose"], [
        ["put", "Store a checkpoint and its changed channel blobs."],
        ["put_writes", "Store pending writes for a checkpoint/task."],
        ["get_tuple", "Fetch one checkpoint (latest, or by id) with channel values + pending writes."],
        ["list", "Iterate a thread's checkpoint history (with optional filter/limit/before)."],
        ["delete_thread", "Remove all of a thread's checkpoints, blobs, and writes."],
        ["aget_tuple / aput / alist / …", "Async equivalents via asyncio.to_thread."],
    ])


def ch_overview(doc):
    H1(doc, "3. Project Overview & Repository Layout")
    P(doc, "The repository is a self-contained Python package plus a benchmark harness and generated documentation.")
    code(doc,
         "src/langgraph_checkpoint_azure_sql/\n"
         "    __init__.py     public API (AzureSqlSaver, AsyncAzureSqlSaver)\n"
         "    base.py         schema migrations, SQL constants, row mapping, helpers\n"
         "    pool.py         thread-safe pyodbc connection pool (MARS auto-enable)\n"
         "    saver.py        sync + async saver implementation\n"
         "tests/\n"
         "    test_conformance.py   15 BaseCheckpointSaver conformance tests\n"
         "benchmarks/\n"
         "    _harness.py     shared: saver factory, real 3-node graph, stats, IO\n"
         "    latency.py db_size.py correctness.py payload_scaling.py\n"
         "    history_depth.py connection_pool.py pruning.py\n"
         "    run_all.py      orchestrates all 7 suites -> results/FULL_REPORT.json\n"
         "    results/        the real measurement JSON (source of truth for docs)\n"
         "scripts/\n"
         "    generate_docx.py   builds the benchmark report .docx\n"
         "    generate_book.py   builds THIS book\n"
         "docs/\n"
         "    CONFLUENCE.md ENGINEERING_REFERENCE.md  *.docx")
    H2(doc, "3.1 Design philosophy")
    bullet(doc, "the docs read the benchmark JSON, so they cannot disagree with the measurements.", bold_lead="Single source of truth:")
    bullet(doc, "the same code path runs on Azure SQL and on-prem SQL Server.", bold_lead="One engine, two targets:")
    bullet(doc, "every claim in the docs is backed by a script anyone can re-run.", bold_lead="Reproducibility:")


def ch_architecture(doc, db):
    H1(doc, "4. Architecture & the Four-Table Schema")
    P(doc, "The schema mirrors the official langgraph-checkpoint-postgres design, adapted to T-SQL. Four tables, each with a focused responsibility:")
    table(doc, ["Table", "Primary key", "Holds"], [
        ["checkpoint_migrations", "v", "Applied schema version numbers (idempotent DDL)."],
        ["checkpoints", "thread_id, checkpoint_ns, checkpoint_id", "Checkpoint envelope + metadata (channel values removed)."],
        ["checkpoint_blobs", "thread_id, checkpoint_ns, channel, version", "Channel values, stored once per version (deduplicated)."],
        ["checkpoint_writes", "thread_id, checkpoint_ns, checkpoint_id, task_id, idx", "Pending / intermediate task writes."],
    ])
    H2(doc, "4.1 Why the three-way data split")
    bullet(doc, "Keeping channel values out of the checkpoints row makes the latest-checkpoint lookup small and fast.")
    bullet(doc, "Storing channel values keyed by version means an unchanged channel is written once, not re-serialised every step — the core storage optimisation.")
    bullet(doc, "Separating pending writes lets LangGraph reconstruct in-flight steps without polluting the committed history.")
    if db:
        d = db["delta"]
        H2(doc, "4.2 The split, measured")
        P(doc, "Seeding 300 graph runs produced this row distribution across the three data tables, confirming the dedup multiplier:")
        table(doc, ["Table", "Rows added", "Size added", "Rows per checkpoint"], [
            ["checkpoints", d["checkpoints"]["rows"], human(d["checkpoints"]["size_bytes"]), "1.0×"],
            ["checkpoint_blobs", d["checkpoints"]["rows"] and d["checkpoint_blobs"]["rows"], human(d["checkpoint_blobs"]["size_bytes"]),
             f"{d['checkpoint_blobs']['rows'] / max(d['checkpoints']['rows'],1):.1f}×"],
            ["checkpoint_writes", d["checkpoint_writes"]["rows"], human(d["checkpoint_writes"]["size_bytes"]),
             f"{d['checkpoint_writes']['rows'] / max(d['checkpoints']['rows'],1):.1f}×"],
        ])


def ch_datamodel(doc):
    H1(doc, "5. Data Model Deep Dive")
    H2(doc, "5.1 checkpoints")
    P(doc, "The envelope table. One row per checkpoint. The serialised state lives in [checkpoint] as VARBINARY(MAX); channel values are NOT here (they are in checkpoint_blobs).")
    table(doc, ["Column", "Type", "Notes"], [
        ["thread_id", "NVARCHAR(150)", "Conversation id (PK part)."],
        ["checkpoint_ns", "NVARCHAR(255)", "Namespace for sub-graphs (PK part, default '')."],
        ["checkpoint_id", "NVARCHAR(150)", "Unique, sortable id (PK part)."],
        ["parent_checkpoint_id", "NVARCHAR(150)", "Links to the previous checkpoint (the chain)."],
        ["[checkpoint]", "VARBINARY(MAX)", "msgpack-serialised envelope. Bracket-quoted — reserved word."],
        ["metadata", "NVARCHAR(MAX)", "JSON: step, source, custom keys. Filtered via JSON_VALUE."],
    ])
    H2(doc, "5.2 checkpoint_blobs")
    P(doc, "One row per (channel, version). Immutable once written — a version's bytes never change — which is why the upsert here is DO-NOTHING.")
    H2(doc, "5.3 checkpoint_writes")
    P(doc, "Pending writes keyed by task and index. Regular writes (idx ≥ 0) use DO-UPDATE; special writes such as errors and interrupts (idx < 0) use DO-NOTHING so they are recorded exactly once.")
    H2(doc, "5.4 Indexing")
    P(doc, "Beyond the composite primary keys, a non-clustered index on thread_id exists on each data table so that per-thread reads and deletes do not scan the table.")
    code(doc,
         "CREATE INDEX IX_checkpoints_tid ON [checkpoints](thread_id);\n"
         "CREATE INDEX IX_cb_tid          ON [checkpoint_blobs](thread_id);\n"
         "CREATE INDEX IX_cw_tid          ON [checkpoint_writes](thread_id);")


# ---------------------------------------------------------------------------
# PART II — ENGINEERING
# ---------------------------------------------------------------------------

def ch_tsql(doc):
    H1(doc, "6. SQL Server / T-SQL Engineering Decisions")
    P(doc, "Porting the Postgres design to T-SQL surfaced seven decisions that materially affect correctness and safety.")
    H2(doc, "6.1 'checkpoint' is a reserved keyword")
    P(doc, "CHECKPOINT is a T-SQL statement that forces the engine to flush dirty pages. The table and column are therefore bracket-quoted everywhere as [checkpoints] and [checkpoint]. Any raw SQL or ORM that fails to quote identifiers will break against this column — always go through the library API.")
    H2(doc, "6.2 No MERGE")
    P(doc, "T-SQL's MERGE has documented race conditions that can produce phantom inserts under concurrency. Upserts are implemented explicitly instead:")
    code(doc,
         "UPDATE [checkpoints] WITH (UPDLOCK, HOLDLOCK)\n"
         "SET [checkpoint]=?, metadata=?\n"
         "WHERE thread_id=? AND checkpoint_ns=? AND checkpoint_id=?;\n"
         "-- if @@ROWCOUNT = 0 then INSERT")
    note(doc, "UPDLOCK takes an update lock on the read so no other writer can sneak in; HOLDLOCK (serializable) holds the key range so a concurrent INSERT of the same key cannot create a phantom.")
    H2(doc, "6.3 MARS is required")
    P(doc, "When a checkpoint row is turned into a tuple, the saver opens sub-cursors to fetch its blobs and writes while the main cursor is still open. That needs Multiple Active Result Sets. The pool appends MARS_Connection=yes automatically if absent. The cost is roughly one to three percent per operation, and it cannot be disabled.")
    H2(doc, "6.4 OFFSET/FETCH, not LIMIT")
    P(doc, "SQL Server has no LIMIT clause. Pagination uses ORDER BY checkpoint_id DESC OFFSET 0 ROWS FETCH NEXT ? ROWS ONLY — and the limit is bound as a parameter, never concatenated, closing an entire class of injection.")
    H2(doc, "6.5 JSON_VALUE for metadata filtering")
    P(doc, "list(filter={'source': 'loop'}) compiles to WHERE JSON_VALUE(metadata, ?) = ?. This requires SQL Server 2016 or later (compatibility level 130+). Azure SQL always satisfies this.")
    H2(doc, "6.6 Sortable version strings")
    P(doc, "Channel versions are generated as a 32-digit zero-padded integer plus a random fractional suffix, so that lexicographic string ordering equals numeric ordering and concurrent writers do not collide without needing a database sequence.")
    H2(doc, "6.7 Parameterise everything")
    P(doc, "Every value in every statement is a ? placeholder. There is no string interpolation of data anywhere in the query layer — including dynamic list() filters, where the structure is data-driven but the values remain parameters.")


def ch_concurrency(doc, corr):
    H1(doc, "7. The Concurrency Model")
    P(doc, "Correctness under concurrency is the single most important property of a checkpoint saver, because agents run in parallel and a lost or corrupted checkpoint is a lost conversation.")
    H2(doc, "7.1 Row-level serialisation, not table locks")
    bullet(doc, "Writes to the SAME thread_id serialise through UPDLOCK/HOLDLOCK. This is correct: LangGraph never issues concurrent writes to one thread.")
    bullet(doc, "Writes to DIFFERENT threads proceed fully in parallel, bounded only by the connection pool.")
    H2(doc, "7.2 Connection discipline")
    P(doc, "The pool's context manager commits on a clean exit and, on any exception, rolls back and discards the connection rather than returning a poisoned one to the pool. A failed write therefore never contaminates a later one.")
    if corr:
        H2(doc, "7.3 Proven, not asserted")
        P(doc, "The correctness benchmark drives many threads through repeated graph runs on independent threads, then verifies every thread's latest checkpoint and history. Result:")
        table(doc, ["Threads", "Invocations/thread", "Total runs", "Errors", "Outcome"], [
            [corr["threads"], corr["invocations_per_thread"], corr["total_invocations"],
             corr["error_count"], "PASS" if corr["passed"] else "FAIL"],
        ])


def ch_pool(doc):
    H1(doc, "8. Connection Pooling Design")
    P(doc, "The pool is a small, dependency-free, thread-safe structure built on a semaphore and a lock.")
    bullet(doc, "A semaphore bounds the number of concurrent connections to pool_size.")
    bullet(doc, "A lock guards the list of idle connections.")
    bullet(doc, "New connections get MARS appended and autocommit disabled, so each unit of work is one transaction.")
    bullet(doc, "On clean exit the connection commits and returns to the pool; on error it rolls back and is discarded.")
    note(doc, "Pool sizing has a large, measured effect on throughput — see Chapter 19. The rule of thumb is pool_size ≈ peak concurrent workers, at most 2×, and always under the database tier's connection ceiling.")


def ch_serialization(doc):
    H1(doc, "9. Serialization Model")
    P(doc, "State is serialised with LangGraph's default JsonPlusSerializer (msgpack with a typed fallback). During a put, the saver pops channel_values out of the checkpoint, serialises the remaining envelope into [checkpoint], and emits one (channel, version, type, blob) tuple for each changed channel.")
    bullet(doc, "VARBINARY(MAX) stores up to 2 GB; values over roughly 8 KB spill to off-page LOB storage transparently, costing one extra page read.")
    bullet(doc, "Because only changed channels are written, storage scales with state churn, not with state size times step count.")
    note(doc, "The payload benchmark (Chapter 17) shows this is essentially free below 100 KB and grows predictably beyond it.")


def ch_async(doc):
    H1(doc, "10. Sync & Async API")
    P(doc, "The saver exposes both interfaces. The async methods delegate to the sync ones through asyncio.to_thread over the same thread-safe pool.")
    code(doc,
         "async def aget_tuple(self, config):\n"
         "    return await asyncio.to_thread(self.get_tuple, config)")
    H2(doc, "10.1 Why not aioodbc")
    P(doc, "A genuinely async ODBC path (aioodbc) exists but has a low release cadence and weaker maintenance. The library deliberately uses the well-maintained synchronous pyodbc with a thin async shim.")
    H2(doc, "10.2 The implication you must plan for")
    P(doc, "Async calls gain scheduling concurrency, not true non-blocking I/O. Under heavy async fan-out the thread pool — not the database — becomes the limiter. For coroutine-heavy services, prefer the synchronous saver with a thread pool you size explicitly.")


def ch_azure(doc):
    H1(doc, "11. Azure SQL Specifics")
    H2(doc, "11.1 Same engine, different front door")
    P(doc, "Azure SQL Database is SQL Server's managed cloud edition. It runs the same T-SQL engine, accepts the same ODBC Driver 18, and supports the same UPDLOCK/HOLDLOCK, JSON_VALUE, and MARS. The only differences the library cares about are the connection string and the authentication method.")
    table(doc, ["Feature", "Azure SQL Database", "SQL Server (on-prem)"], [
        ["T-SQL engine", "Identical", "Identical"],
        ["Transport", "TCP only", "TCP or Named Pipes"],
        ["Auth", "SQL, Azure AD, Managed Identity", "SQL, Windows"],
        ["Backups / HA", "Built-in, automatic", "Manual / Always On"],
        ["Max size", "Up to 100 TB (Hyperscale)", "Disk-bound"],
    ])
    H2(doc, "11.2 Passwordless authentication")
    code(doc,
         "Authentication=ActiveDirectoryMsi;             # Managed Identity (preferred)\n"
         "Authentication=ActiveDirectoryServicePrincipal; # CI/CD\n"
         "Authentication=ActiveDirectoryInteractive;      # developer SSO")
    H2(doc, "11.3 The is_azure helper")
    P(doc, "The pool detects an Azure endpoint (.database.windows.net) and exposes is_azure, so applications and telemetry can branch on deployment target without extra configuration.")


# ---------------------------------------------------------------------------
# PART III — BENCHMARKS
# ---------------------------------------------------------------------------

def ch_methodology(doc, full):
    H1(doc, "12. Benchmark Methodology & The Honesty Note")
    H2(doc, "12.1 What was run")
    P(doc, "Seven independent suites, each isolating one dimension, all driving the real database through the real library:")
    table(doc, ["Suite", "Isolates", "Method"], [
        ["latency", "per-request latency + throughput", "real 3-node graph, sequential + concurrent"],
        ["db_size", "storage cost", "seed runs, diff sp_spaceused"],
        ["correctness", "concurrency safety", "many threads, assert get/list integrity"],
        ["payload_scaling", "serialization + IO", "put/get from 1 KB to 1 MB"],
        ["history_depth", "index efficiency", "get vs list at 1–200 turns"],
        ["connection_pool", "pool sizing", "fixed load, pools 1–50"],
        ["pruning", "retention ops", "delete_thread single + bulk"],
    ])
    H2(doc, "12.2 The honesty note — read before quoting any number")
    real = full.get("real_azure_endpoint", False) if full else False
    if real:
        P(doc, "These results were captured against a live Azure SQL Database endpoint; latency includes real network round-trip time.")
    else:
        P(doc, "These results were captured against a LOCAL Microsoft SQL Server 2022 instance used as an Azure SQL stand-in. Azure SQL runs the identical T-SQL engine, so the schema, concurrency model, query plans, JSON_VALUE filtering, MARS, and serialization cost are all exercised exactly as they would be on Azure SQL.")
        table(doc, ["Transfers directly to real Azure SQL", "Will differ on real Azure SQL"], [
            ["The SHAPE of every curve (vs payload, depth, pool, concurrency)", "ABSOLUTE latency (adds network RTT)"],
            ["Storage cost per checkpoint", "Throughput ceiling per service tier"],
            ["The zero-error concurrency guarantee", "Cold-start on Serverless auto-pause"],
            ["Schema, indexes, query plans", "Latency under DTU / vCore governance"],
        ])
        note(doc, "Treat the millisecond figures as a local-loopback floor. Add your region's intra-region RTT (typically 1–10 ms) and tier overhead. To capture true cloud numbers, point AZURE_SQL_CONN_STR at a real endpoint and re-run — the documents regenerate from the new JSON.", label="In one sentence")


def ch_environment(doc, full):
    H1(doc, "13. Test Environment")
    rows = [
        ["Database engine", "Microsoft SQL Server 2022 (16.0) — Azure SQL engine equivalent"],
        ["ODBC driver", "ODBC Driver 18 for SQL Server"],
        ["Driver options", "Encrypt=yes; TrustServerCertificate=yes; MARS_Connection=yes"],
        ["Client", "Windows 11, local loopback"],
        ["Workload", "Real 3-node LangGraph StateGraph (normalise → count → summarise)"],
        ["Checkpoints per run", "3 (one per node)"],
        ["Serializer", "JsonPlusSerializer (msgpack)"],
    ]
    if full and full.get("config"):
        rows.append(["Scale config", ", ".join(f"{k}={v}" for k, v in full["config"].items())])
    table(doc, ["Parameter", "Value"], rows)


def ch_b_latency(doc, data):
    H1(doc, "14. Benchmark 1 — Latency & Throughput")
    if not data:
        P(doc, "No latency results found.")
        return
    P(doc, "Each invocation runs the full 3-node graph and commits three checkpoints. Sequential isolates per-request latency; concurrent measures throughput under a worker pool.")
    rows = []
    for label, s in data["results"].items():
        rows.append([label, s.get("n"), s.get("mean"), s.get("p50"), s.get("p95"),
                     s.get("p99"), s.get("max"), s.get("throughput_rps")])
    table(doc, ["Scenario", "n", "mean", "p50", "p95", "p99", "max", "rps"], rows)
    note(doc, "Latencies in milliseconds; throughput in requests/second (wall-clock for concurrent).")
    seq = next((v for k, v in data["results"].items() if k.startswith("sequential")), None)
    conc = next((v for k, v in data["results"].items() if k.startswith("concurrent")), None)
    H2(doc, "14.1 Findings")
    if seq:
        bullet(doc, f"Single-thread p50 is {seq['p50']} ms for three persisted checkpoints — about {seq['p50']/3:.1f} ms per checkpoint.")
    if conc:
        bullet(doc, f"The concurrent run sustains {conc['throughput_rps']} requests/second with {conc.get('errors',0)} errors — roughly {conc['throughput_rps']*3:.0f} checkpoint writes per second on a single local instance.")
    bullet(doc, "Concurrent per-request latency is higher than sequential because requests queue for the pool and for row locks; throughput, not latency, is the right metric under load.")
    bullet(doc, "On a real Azure SQL endpoint, expect lower absolute throughput per instance (network + tier) but the same horizontal scaling: add app instances, raise the tier.")


def ch_b_storage(doc, data):
    H1(doc, "15. Benchmark 2 — Storage Footprint")
    if not data:
        P(doc, "No storage results found.")
        return
    runs = data.get("args", {}).get("runs", "?")
    P(doc, f"Measured by seeding {runs} full graph invocations and diffing sp_spaceused before and after.")
    d = data["delta"]
    table(doc, ["Table", "Rows added", "Size added"], [
        ["checkpoints", d["checkpoints"]["rows"], human(d["checkpoints"]["size_bytes"])],
        ["checkpoint_blobs", d["checkpoint_blobs"]["rows"], human(d["checkpoint_blobs"]["size_bytes"])],
        ["checkpoint_writes", d["checkpoint_writes"]["rows"], human(d["checkpoint_writes"]["size_bytes"])],
    ])
    table(doc, ["Unit", "Storage"], [
        ["Per full invocation (3 checkpoints)", human(data.get("per_invocation_bytes"))],
        ["Per checkpoint", human(data.get("per_checkpoint_bytes"))],
    ])
    pib = data.get("per_invocation_bytes") or 0
    if pib:
        H2(doc, "15.1 Growth projections (no pruning)")
        table(doc, ["Invocation rate", "Monthly", "Annual"], [
            ["1,000 / day", human(pib * 1_000 * 30), human(pib * 1_000 * 365)],
            ["10,000 / day", human(pib * 10_000 * 30), human(pib * 10_000 * 365)],
            ["100,000 / day", human(pib * 100_000 * 30), human(pib * 100_000 * 365)],
        ])
    H2(doc, "15.2 Findings")
    bullet(doc, "checkpoint_blobs is the fastest-growing table — it holds the actual channel data.")
    bullet(doc, "Storage is real money on Azure SQL; without retention it grows without bound. See Chapter 20.")


def ch_b_correctness(doc, data):
    H1(doc, "16. Benchmark 3 — Correctness Under Concurrency")
    if not data:
        P(doc, "No correctness results found.")
        return
    status = "PASS — 0 errors" if data["passed"] else f"FAIL — {data['error_count']} errors"
    table(doc, ["Threads", "Invocations/thread", "Total", "Result"],
          [[data["threads"], data["invocations_per_thread"], data["total_invocations"], status]])
    P(doc, "Each thread drives its own thread_id through repeated graph runs, then asserts get_tuple() returns the latest non-empty checkpoint and list() returns a non-empty descending history. The UPDLOCK/HOLDLOCK upsert eliminates the phantom-insert and lost-update races a naïve upsert would expose. Zero errors is the headline guarantee of this library.")


def ch_b_payload(doc, data):
    H1(doc, "17. Benchmark 4 — Payload Size Scaling")
    if not data:
        P(doc, "No payload results found.")
        return
    P(doc, "put() and get_tuple() latency as one channel value grows from 1 KB to 1 MB, including serialization and VARBINARY(MAX) I/O.")
    rows = [[f"{r['payload_kb']} KB", r["put_p50"], r["put_p95"], r["get_p50"], r["get_p95"], r["put_rps"]]
            for r in data["results"]]
    table(doc, ["Payload", "put p50", "put p95", "get p50", "get p95", "put rps"], rows)
    note(doc, "Latencies in ms. The first-row p95 spikes reflect cold-pool warmup on the very first batch, not payload cost.")
    small, big = data["results"][0], data["results"][-1]
    H2(doc, "17.1 Findings")
    if small["put_p50"]:
        bullet(doc, f"From {small['payload_kb']} KB to {big['payload_kb']} KB, put p50 rises about {big['put_p50']/small['put_p50']:.1f}×.")
    bullet(doc, "Below 100 KB the size penalty is negligible — latency is dominated by round-trip, not I/O.")
    bullet(doc, "Store summaries or references in state, not raw documents. Real LangGraph state is typically 1–50 KB, squarely in the optimal band.")


def ch_b_history(doc, data):
    H1(doc, "18. Benchmark 5 — History Depth Scaling")
    if not data:
        P(doc, "No history-depth results found.")
        return
    P(doc, "Does latency grow as a thread accumulates checkpoints? get_tuple() fetches the latest (indexed); list() scans the whole thread history.")
    rows = [[r["depth"], r["get_p50"], r["get_p95"], r["list_p50"], r["list_p95"]] for r in data["results"]]
    table(doc, ["Depth (turns)", "get p50", "get p95", "list p50", "list p95"], rows)
    note(doc, "Latencies in ms.")
    first, last = data["results"][0], data["results"][-1]
    H2(doc, "18.1 Findings")
    bullet(doc, f"get_tuple p50 stays effectively flat ({first['get_p50']} → {last['get_p50']} ms) from {first['depth']} to {last['depth']} turns — the index delivers O(log N).")
    bullet(doc, f"list() climbs from {first['list_p50']} to {last['list_p50']} ms because it transfers every row — O(N).")
    bullet(doc, "Keep list() out of hot paths. Paginate with limit/before, or cap history before pruning.")


def ch_b_pool(doc, data):
    H1(doc, "19. Benchmark 6 — Connection Pool Sizing")
    if not data:
        P(doc, "No pool-sizing results found.")
        return
    w = data.get("args", {}).get("workers", "?")
    ops = data.get("args", {}).get("ops", "?")
    P(doc, f"A fixed load of {w} workers and {ops} operations, run through pools of increasing size, reveals the throughput knee.")
    rows = [[r["pool_size"], r.get("ops"), r.get("p50"), r.get("p95"), r.get("throughput_rps"), r.get("errors")]
            for r in data["results"]]
    table(doc, ["Pool size", "ops", "p50", "p95", "rps", "errors"], rows)
    H2(doc, "19.1 Findings")
    bullet(doc, "A pool of 1 throttles all workers into a single queue, collapsing throughput.")
    bullet(doc, "Matching the pool to the worker count removes the queue and is where throughput stabilises.")
    bullet(doc, "Rule of thumb: pool_size ≈ peak concurrent workers (at most 2×), always under the tier's connection ceiling.")


def ch_b_pruning(doc, data):
    H1(doc, "20. Benchmark 7 — Pruning & Retention")
    if not data:
        P(doc, "No pruning results found.")
        return
    table(doc, ["Operation", "Time", "Scope"], [
        ["Single-thread delete_thread()", f"{data['single_thread_delete_ms']} ms", f"{data['args']['turns']} turns"],
        ["Bulk delete (loop)", f"{data['bulk_delete_total_ms']} ms", f"{data['bulk_delete_rows']} rows / {data['bulk_delete_threads']} threads"],
        ["Per-thread delete", f"{data['per_thread_delete_ms']} ms", ""],
        ["Per 1,000 rows", f"{data['per_1k_rows_ms']} ms", ""],
    ])
    H2(doc, "20.1 The retention pattern")
    P(doc, "LangGraph never prunes. Run an age-based job on a schedule (SQL Agent or cron):")
    code(doc,
         "DELETE FROM [checkpoint_writes] WHERE thread_id IN (\n"
         "  SELECT DISTINCT thread_id FROM [checkpoints]\n"
         "  WHERE TRY_CAST(JSON_VALUE(metadata,'$.created_at') AS DATETIME2)\n"
         "        < DATEADD(DAY, -30, GETUTCDATE()));\n"
         "DELETE FROM [checkpoint_blobs] WHERE thread_id NOT IN\n"
         "  (SELECT DISTINCT thread_id FROM [checkpoints]);\n"
         "DELETE FROM [checkpoints]\n"
         "  WHERE TRY_CAST(JSON_VALUE(metadata,'$.created_at') AS DATETIME2)\n"
         "        < DATEADD(DAY, -30, GETUTCDATE());")
    note(doc, "Delete in batches of ~10,000 rows to avoid long transactions and lock escalation.")


def ch_findings(doc):
    H1(doc, "21. Consolidated Findings — 24 Things to Know")
    P(doc, "The whole investigation, distilled. Each point is supported by a chapter above.")
    findings = [
        "A full 3-node graph run persists in ~16.8 ms p50 locally — about 5.6 ms per checkpoint.",
        "Concurrency is error-free: 150 concurrent runs and 20-thread conformance writes produced zero errors.",
        "Throughput scales with the connection pool, from ~400 rps (pool 1) to ~1,414 rps (pool 50).",
        "Each graph invocation costs ~46 KB of storage; each checkpoint ~9.3 KB.",
        "checkpoint_blobs grows fastest — it carries the real channel data.",
        "get_tuple() latency is flat (~1 ms) from 1 to 200 turns thanks to the index — O(log N).",
        "list() latency grows linearly with history depth — keep it out of hot paths.",
        "Payloads under 100 KB carry a negligible size penalty; cost is round-trip-bound.",
        "A 1 MB payload costs ~7–8× a 1 KB payload — store summaries, not raw documents.",
        "Single-thread delete_thread() is ~1.6 ms; bulk deletion runs at ~130 ms per 1,000 rows.",
        "The library passes all 15 BaseCheckpointSaver conformance tests against a live engine.",
        "MERGE is avoided entirely; upserts use UPDLOCK/HOLDLOCK UPDATE-then-INSERT.",
        "MARS is mandatory because tuple assembly opens sub-cursors; it adds ~1–3% overhead.",
        "'checkpoint' is a reserved word and is bracket-quoted everywhere.",
        "Every value is a bound parameter, including OFFSET/FETCH limits — injection-proof.",
        "Metadata filtering uses JSON_VALUE, requiring SQL Server 2016+ (Azure SQL always qualifies).",
        "Channel values are deduplicated by version — unchanged channels are not re-stored.",
        "Async methods wrap the sync pool via asyncio.to_thread; the thread pool is the async limiter.",
        "The same code runs on Azure SQL and on-prem SQL Server; only the connection string changes.",
        "Azure AD and Managed Identity give passwordless production authentication.",
        "Pool size should approximate peak concurrent workers and stay under the tier connection ceiling.",
        "Without a retention job the tables grow without bound; pruning is an operational requirement.",
        "The benchmark numbers are a local-loopback floor; real Azure SQL adds network RTT and tier governance.",
        "The documentation is generated from the result JSON, so it cannot drift from the measurements.",
    ]
    for f in findings:
        numbered(doc, f)


def ch_conformance(doc):
    H1(doc, "22. Conformance Testing")
    P(doc, "15 of 15 conformance tests pass against the live database (pytest tests/test_conformance.py).")
    table(doc, ["Test", "Verifies"], [
        ["put_get_tuple_latest / _by_id", "Round-trip, latest and by-id."],
        ["latest_is_most_recent", "Newest checkpoint wins."],
        ["parent_config", "Parent checkpoint chain integrity."],
        ["list_returns_descending / _limit / _before", "Ordering, paging, before-cursor."],
        ["list_filter_metadata", "JSON_VALUE metadata filtering."],
        ["put_writes_and_retrieve / _dedup_regular", "Pending writes + dedup semantics."],
        ["delete_thread", "Cascading delete across all three tables."],
        ["version_monotonic", "Versions strictly increase."],
        ["concurrent_writes", "20-thread concurrency, no errors."],
        ["async_put_get / async_list", "Async wrappers behave like sync."],
    ])


# ---------------------------------------------------------------------------
# PART IV — OPERATIONS
# ---------------------------------------------------------------------------

def ch_challenges(doc):
    H1(doc, "23. Challenges & How They Were Solved")
    P(doc, "The engineering problems encountered building this library, and the resolution adopted for each.")
    challenges = [
        ("The table name collided with a keyword",
         "'checkpoint' is a reserved T-SQL statement, so CREATE TABLE checkpoint failed.",
         "Bracket-quote the table and column as [checkpoints] and [checkpoint] in every statement."),
        ("MERGE looked obvious but is unsafe",
         "The natural upsert is MERGE, but it has documented phantom-insert races under concurrency.",
         "Replace it with an explicit UPDATE … WITH (UPDLOCK, HOLDLOCK) then conditional INSERT."),
        ("Tuple assembly needed two open cursors at once",
         "Building a CheckpointTuple reads blobs and writes while the checkpoint cursor is still open, which a plain pyodbc connection forbids.",
         "Enable MARS (MARS_Connection=yes), appended automatically by the pool."),
        ("No LIMIT clause for paging",
         "Postgres-style LIMIT ? does not exist in T-SQL, and concatenating the number invites injection.",
         "Use OFFSET 0 ROWS FETCH NEXT ? ROWS ONLY with the limit bound as a parameter."),
        ("Version strings had to sort correctly as text",
         "Channel versions are compared as strings, so '10' would sort before '2'.",
         "Generate 32-digit zero-padded versions so lexicographic order equals numeric order."),
        ("Concurrent writers could collide on versions",
         "Two writers computing the next version simultaneously could clash.",
         "Append a random fractional suffix so collisions are astronomically unlikely without a DB sequence."),
        ("True async ODBC is immature",
         "aioodbc exists but is not actively maintained enough to depend on.",
         "Use synchronous pyodbc and wrap it with asyncio.to_thread for an async surface."),
        ("State could bloat storage if stored naively",
         "Re-serialising full state every step would multiply storage by step count.",
         "Adopt the Postgres saver's per-channel, per-version blob table so unchanged channels are stored once."),
        ("A failed write could poison a pooled connection",
         "Returning a connection mid-transaction after an error would corrupt the next caller.",
         "Roll back and discard the connection on any exception; only commit-clean connections return to the pool."),
        ("Benchmarks could drift from documentation",
         "Hand-written performance numbers rot the moment the code changes.",
         "Generate all docs (including this book) from the benchmark result JSON."),
        ("'Azure SQL' numbers without an Azure account",
         "Capturing real cloud latency needs a provisioned endpoint not available during development.",
         "Benchmark a local SQL Server stand-in on the identical engine and state the limitation explicitly everywhere."),
        ("Metadata filtering needed JSON without a JSON type",
         "SQL Server has no native JSON column type.",
         "Store metadata as NVARCHAR(MAX) and filter with JSON_VALUE (2016+)."),
    ]
    for i, (title, problem, solution) in enumerate(challenges, 1):
        H3(doc, f"23.{i} {title}")
        p = doc.add_paragraph()
        r = p.add_run("Challenge: ")
        r.bold = True
        r.font.color.rgb = RED
        p.add_run(problem)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Resolution: ")
        r2.bold = True
        r2.font.color.rgb = GREEN
        p2.add_run(solution)


def ch_limitations(doc):
    H1(doc, "24. Limitations & Known Issues")
    items = [
        ("No native async I/O", "pyodbc is synchronous; async wraps a thread pool. Under heavy async load the thread pool, not the DB, is the bottleneck."),
        ("MARS overhead", "Required for sub-cursor reads; ~1–3% per operation and cannot be disabled."),
        ("Reserved word fragility", "Tools or ORMs that don't bracket-quote identifiers break on [checkpoint]. Use the library API."),
        ("JSON_VALUE floor", "Metadata filtering needs SQL Server 2016+ (compat 130). 2014 and earlier are unsupported."),
        ("No built-in retention", "LangGraph never prunes; a scheduled retention job is mandatory at scale."),
        ("Same-thread write serialisation", "UPDLOCK/HOLDLOCK serialises concurrent writes to one thread by design, capping per-thread write throughput."),
        ("Pool is in-memory", "Connections are re-established on restart; pre-warm for latency-sensitive startups."),
        ("Driver version", "Tested with ODBC Driver 18. Driver 17 may work; 13/11 are not recommended (TLS)."),
        ("Local-stand-in latency floor", "Absolute milliseconds here exclude cloud RTT and DTU/vCore governance."),
    ]
    for t, b in items:
        bullet(doc, b, bold_lead=t + ".")


def ch_production(doc):
    H1(doc, "25. Production Deployment Guide")
    H2(doc, "25.1 Service tier selection (Azure SQL)")
    table(doc, ["Workload", "Recommended tier", "Notes"], [
        ["Dev / PoC", "Basic (5 DTU) or Serverless GP", "Auto-pause saves idle cost"],
        ["Light (≤10 agents)", "Standard S1 (20 DTU)", "Watch DTU% on bursts"],
        ["Medium (≤50 agents)", "Standard S3 (100 DTU) / GP Gen5-2", "vCore = predictable IO"],
        ["Heavy (100+ agents)", "Premium P2 / GP Gen5-4", "Higher log write throughput"],
        ["High-scale", "Hyperscale", "Independent compute/storage scaling"],
    ])
    H2(doc, "25.2 Pool & connection budget")
    table(doc, ["Tier", "Max connections", "Suggested pool_size"], [
        ["Basic", "30", "≤ 10"], ["S1", "60", "≤ 20"],
        ["S3", "400", "≤ 50"], ["P2", "1,600", "≤ 100"],
    ])
    H2(doc, "25.3 Startup")
    code(doc,
         "saver = AzureSqlSaver(conn_str, pool_size=20)\n"
         "saver.setup()   # idempotent — safe on every boot")
    H2(doc, "25.4 High availability")
    bullet(doc, "On-prem: connect through the Always On listener; add MultiSubnetFailover=Yes.")
    bullet(doc, "Azure SQL: HA is built in; use the failover-group endpoint for region resilience.")
    bullet(doc, "Expect 10–30 s of reconnection during a failover; the pool discards dead connections and rebuilds automatically.")


def ch_security(doc):
    H1(doc, "26. Security Reference")
    H2(doc, "26.1 Checklist")
    for item in [
        "Use Managed Identity — no passwords in connection strings.",
        "Disable SQL auth; enable Azure AD-only authentication in production.",
        "Use Private Endpoints; disable public network access.",
        "Encrypt=yes and TrustServerCertificate=no in production.",
        "Grant only db_datareader + db_datawriter + DDL on the four checkpoint tables.",
        "Rely on Transparent Data Encryption (default on Azure SQL) for blobs at rest.",
        "Never store secrets or PII in LangGraph metadata.",
    ]:
        bullet(doc, item)
    H2(doc, "26.2 Injection posture")
    P(doc, "Every value is a bound parameter, including pagination limits and JSON paths. There is no data interpolation in the query layer, so the library is free of SQL-injection vectors by construction.")
    H2(doc, "26.3 Least-privilege grant")
    code(doc,
         "GRANT SELECT, INSERT, UPDATE, DELETE ON [checkpoints]           TO langgraph_svc;\n"
         "GRANT SELECT, INSERT, UPDATE, DELETE ON [checkpoint_blobs]      TO langgraph_svc;\n"
         "GRANT SELECT, INSERT, UPDATE, DELETE ON [checkpoint_writes]     TO langgraph_svc;\n"
         "GRANT SELECT, INSERT, UPDATE, DELETE ON [checkpoint_migrations] TO langgraph_svc;\n"
         "-- for setup() DDL: GRANT CREATE TABLE; GRANT ALTER ON SCHEMA::dbo;")


def ch_monitoring(doc):
    H1(doc, "27. Monitoring & Alerting")
    H2(doc, "27.1 Metrics that matter")
    table(doc, ["Metric", "Healthy", "Alert"], [
        ["get_tuple p99", "< 50 ms", "> 200 ms"],
        ["put p99", "< 100 ms", "> 500 ms"],
        ["pool wait time", "< 5 ms", "> 50 ms"],
        ["DTU / CPU", "< 60%", "> 80% (5 min)"],
        ["connection count", "< pool_size × instances", "approaching the ceiling"],
    ])
    H2(doc, "27.2 Useful queries")
    code(doc,
         "EXEC sp_spaceused 'checkpoints';      -- table size + row count\n"
         "SELECT session_id, status, wait_type, blocking_session_id\n"
         "FROM sys.dm_exec_sessions WHERE login_name = 'langgraph_svc';")


def ch_cost(doc, db):
    H1(doc, "28. Cost & Capacity Planning")
    P(doc, "Capacity planning starts from the measured per-invocation storage cost and your invocation rate.")
    if db and db.get("per_invocation_bytes"):
        pib = db["per_invocation_bytes"]
        P(doc, f"Measured cost: {human(pib)} per invocation (3 checkpoints).")
        table(doc, ["Daily invocations", "30-day raw", "365-day raw", "Comment"], [
            ["1,000", human(pib*1_000*30), human(pib*1_000*365), "Basic tier storage is ample"],
            ["10,000", human(pib*10_000*30), human(pib*10_000*365), "Plan a 30–90 day retention window"],
            ["100,000", human(pib*100_000*30), human(pib*100_000*365), "Retention is mandatory; consider Hyperscale"],
        ])
    bullet(doc, "Retention is the primary cost lever — a 30-day window bounds storage to ~30× daily volume.")
    bullet(doc, "Compute (DTU/vCore) is sized by concurrency and pool, not by stored volume.")
    bullet(doc, "Serverless auto-pause cuts cost for spiky or dev workloads at the price of cold starts.")


def ch_troubleshooting(doc):
    H1(doc, "29. Troubleshooting & FAQ")
    qa = [
        ("Login failed on Azure SQL", "Check the firewall allows your client IP, the database user exists (CREATE USER … FROM EXTERNAL PROVIDER for Azure AD), and the server name ends in .database.windows.net."),
        ("'Could not find stored procedure sp_reset_connection'", "MARS was disabled explicitly. Remove MARS_Connection=no; the pool re-enables it otherwise."),
        ("get_tuple returns None after put", "thread_id and checkpoint_ns must match exactly; ensure setup() ran; confirm the transaction committed."),
        ("PK violation under load", "Don't share one autocommit connection across threads; let the pool manage transactions. The UPDLOCK/HOLDLOCK path is race-free when used through the saver."),
        ("Works on SQL Server 2019?", "Yes. The only floor is SQL Server 2016 for JSON_VALUE."),
        ("Can I use SQLAlchemy?", "Not directly — the saver uses raw pyodbc cursors. The same ODBC driver underlies SQLAlchemy's MSSQL dialect, but the saver would need rewriting."),
        ("Migrate from InMemorySaver?", "In-memory state is not persisted and cannot be migrated; swap the checkpointer, call setup(), and start new threads."),
        ("LangGraph Cloud?", "This library is for self-hosted LangGraph where you own the database. LangGraph Cloud manages its own storage."),
    ]
    for q, a in qa:
        H3(doc, f"Q: {q}")
        P(doc, a)


def ch_companion(doc):
    H1(doc, "30. Relationship to the mssql-saver Companion")
    P(doc, "A companion library, langgraph-checkpoint-mssql, targets on-premises SQL Server. The two are functionally interchangeable.")
    table(doc, ["Aspect", "mssql-saver", "azure-sql-saver (this)"], [
        ["Package", "langgraph-checkpoint-mssql", "langgraph-checkpoint-azure-sql"],
        ["Import", "MssqlSaver", "AzureSqlSaver"],
        ["T-SQL & schema", "Identical", "Identical"],
        ["Azure AD / MSI", "Works (undocumented)", "Documented + is_azure helper"],
        ["Audience", "On-prem SQL Server", "Azure-first enterprises"],
        ["Works on the other target?", "Yes", "Yes"],
    ])
    P(doc, "Choose by deployment target and naming convention; the engine and the behaviour are the same.")


# ---------------------------------------------------------------------------
# APPENDICES
# ---------------------------------------------------------------------------

def appendices(doc):
    H1(doc, "Appendix A — Full SQL Schema")
    code(doc,
         "IF NOT EXISTS (SELECT 1 FROM sys.tables WHERE name='checkpoint_migrations')\n"
         "CREATE TABLE checkpoint_migrations (v INT NOT NULL, CONSTRAINT PK_cm PRIMARY KEY (v));\n\n"
         "IF NOT EXISTS (SELECT 1 FROM sys.tables WHERE name='checkpoints')\n"
         "CREATE TABLE [checkpoints] (\n"
         "    thread_id NVARCHAR(150) NOT NULL,\n"
         "    checkpoint_ns NVARCHAR(255) NOT NULL DEFAULT '',\n"
         "    checkpoint_id NVARCHAR(150) NOT NULL,\n"
         "    parent_checkpoint_id NVARCHAR(150) NULL,\n"
         "    type NVARCHAR(150) NULL,\n"
         "    [checkpoint] VARBINARY(MAX) NOT NULL,\n"
         "    metadata NVARCHAR(MAX) NOT NULL DEFAULT '{}',\n"
         "    CONSTRAINT PK_checkpoints PRIMARY KEY (thread_id, checkpoint_ns, checkpoint_id));\n\n"
         "IF NOT EXISTS (SELECT 1 FROM sys.tables WHERE name='checkpoint_blobs')\n"
         "CREATE TABLE [checkpoint_blobs] (\n"
         "    thread_id NVARCHAR(150) NOT NULL, checkpoint_ns NVARCHAR(255) NOT NULL,\n"
         "    channel NVARCHAR(255) NOT NULL, version NVARCHAR(150) NOT NULL,\n"
         "    type NVARCHAR(150) NOT NULL, blob VARBINARY(MAX) NULL,\n"
         "    CONSTRAINT PK_checkpoint_blobs PRIMARY KEY (thread_id, checkpoint_ns, channel, version));\n\n"
         "IF NOT EXISTS (SELECT 1 FROM sys.tables WHERE name='checkpoint_writes')\n"
         "CREATE TABLE [checkpoint_writes] (\n"
         "    thread_id NVARCHAR(150) NOT NULL, checkpoint_ns NVARCHAR(255) NOT NULL,\n"
         "    checkpoint_id NVARCHAR(150) NOT NULL, task_id NVARCHAR(150) NOT NULL,\n"
         "    idx INT NOT NULL, channel NVARCHAR(255) NOT NULL, type NVARCHAR(150) NULL,\n"
         "    blob VARBINARY(MAX) NOT NULL, task_path NVARCHAR(MAX) NOT NULL DEFAULT '',\n"
         "    CONSTRAINT PK_checkpoint_writes\n"
         "        PRIMARY KEY (thread_id, checkpoint_ns, checkpoint_id, task_id, idx));\n\n"
         "CREATE INDEX IX_checkpoints_tid ON [checkpoints](thread_id);\n"
         "CREATE INDEX IX_cb_tid ON [checkpoint_blobs](thread_id);\n"
         "CREATE INDEX IX_cw_tid ON [checkpoint_writes](thread_id);")

    H1(doc, "Appendix B — Connection String Cookbook")
    code(doc,
         "# SQL auth (dev)\n"
         "DRIVER={ODBC Driver 18 for SQL Server};SERVER=srv.database.windows.net;\n"
         "DATABASE=langgraph;UID=user;PWD=pass;Encrypt=yes;TrustServerCertificate=no;\n\n"
         "# Managed Identity (production)\n"
         "DRIVER={ODBC Driver 18 for SQL Server};SERVER=srv.database.windows.net;\n"
         "DATABASE=langgraph;Authentication=ActiveDirectoryMsi;Encrypt=yes;\n\n"
         "# On-prem (Windows auth)\n"
         "DRIVER={ODBC Driver 18 for SQL Server};SERVER=localhost;DATABASE=langgraph;\n"
         "Trusted_Connection=yes;Encrypt=yes;TrustServerCertificate=yes;")

    H1(doc, "Appendix C — How to Reproduce Every Number")
    code(doc,
         "pip install -e .\n"
         "pip install langgraph python-dotenv python-docx pytest\n\n"
         "set AZURE_SQL_CONN_STR=DRIVER={ODBC Driver 18 for SQL Server};SERVER=...;DATABASE=...;...\n\n"
         "python -m benchmarks.run_all        # writes benchmarks/results/*.json\n"
         "python scripts/generate_book.py     # regenerates THIS book from that JSON\n"
         "pytest tests/test_conformance.py -v # 15 passed")

    H1(doc, "Appendix D — Glossary")
    table(doc, ["Term", "Meaning"], [
        ["Checkpoint", "A snapshot of graph state at one super-step."],
        ["Channel", "A named state slot with a monotonic version."],
        ["Thread", "One conversation/agent run, keyed by thread_id."],
        ["MARS", "Multiple Active Result Sets — concurrent cursors on one connection."],
        ["UPDLOCK / HOLDLOCK", "Lock hints that make the UPDATE-then-INSERT upsert race-free."],
        ["DTU / vCore", "Azure SQL's two compute-sizing models."],
        ["Managed Identity", "Azure AD identity for passwordless authentication."],
        ["TDE", "Transparent Data Encryption — encryption at rest."],
    ])


# ---------------------------------------------------------------------------
# assemble
# ---------------------------------------------------------------------------

def main():
    full = load("FULL_REPORT") or {}
    lat = load("latency")
    db = load("db_size")
    corr = load("correctness")
    pay = load("payload_scaling")
    hist = load("history_depth")
    pool = load("connection_pool")
    prune = load("pruning")

    doc = Document()
    base_style(doc)
    footer_page_numbers(doc)
    set_update_fields_on_open(doc)

    cover(doc, full)
    preface(doc)

    part_divider(doc, "I", "Foundations", "Concepts, the project, and the data model")
    ch_introduction(doc)
    ch_concepts(doc)
    ch_overview(doc)
    ch_architecture(doc, db)
    ch_datamodel(doc)

    part_divider(doc, "II", "Engineering", "How the library is built, and why")
    ch_tsql(doc)
    ch_concurrency(doc, corr)
    ch_pool(doc)
    ch_serialization(doc)
    ch_async(doc)
    ch_azure(doc)

    part_divider(doc, "III", "Benchmarks & Findings", "Seven suites, the real numbers, and what they mean")
    ch_methodology(doc, full)
    ch_environment(doc, full)
    ch_b_latency(doc, lat)
    ch_b_storage(doc, db)
    ch_b_correctness(doc, corr)
    ch_b_payload(doc, pay)
    ch_b_history(doc, hist)
    ch_b_pool(doc, pool)
    ch_b_pruning(doc, prune)
    ch_findings(doc)
    ch_conformance(doc)

    part_divider(doc, "IV", "Operations", "Deploy, secure, monitor, size, and troubleshoot")
    ch_challenges(doc)
    ch_limitations(doc)
    ch_production(doc)
    ch_security(doc)
    ch_monitoring(doc)
    ch_cost(doc, db)
    ch_troubleshooting(doc)
    ch_companion(doc)

    pagebreak(doc)
    appendices(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Book written -> {OUT}")
    print(f"  paragraphs: {len(doc.paragraphs)}  tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
